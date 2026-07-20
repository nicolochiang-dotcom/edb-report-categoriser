"""Extract plain text from report files (.docx, .pdf, .txt, .md)."""

from pathlib import Path

SUPPORTED_EXTENSIONS = {".docx", ".pdf", ".txt", ".md"}


def extract_docx(path: Path) -> str:
    """Read a Word document, including text inside tables."""
    from docx import Document

    doc = Document(str(path))
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    # Many call reports are laid out as tables; capture those too.
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n\n".join(parts)


def extract_pdf(path: Path) -> str:
    """Read a text-based PDF page by page."""
    import pdfplumber

    pages = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            pages.append(page.extract_text() or "")
    return "\n\n".join(pages)


def extract_text(path: Path) -> str:
    """Route a file to the right extractor based on its extension."""
    ext = path.suffix.lower()
    if ext == ".docx":
        return extract_docx(path)
    if ext == ".pdf":
        return extract_pdf(path)
    if ext in {".txt", ".md"}:
        return path.read_text(encoding="utf-8", errors="ignore")
    raise ValueError(f"Unsupported file type: {ext}")